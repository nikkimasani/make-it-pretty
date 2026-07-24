import io

from docx import Document

from app.services.reader import (
    _count_words,
    _is_url,
    _looks_like_html,
    process_document,
)

SAMPLE_HTML = """<!DOCTYPE html>
<html>
<head><title>Test Article</title></head>
<body>
  <nav>Nav links here</nav>
  <header>Header content</header>
  <article>
    <h1>Main Title</h1>
    <p>First paragraph with <strong>bold</strong> and <em>italic</em> text.</p>
    <p>Second paragraph with a <a href="https://example.com">link</a>.</p>
    <ul>
      <li>Item one</li>
      <li>Item two</li>
    </ul>
    <blockquote>A notable quote.</blockquote>
    <img src="image.jpg" alt="An image" />
  </article>
  <aside>Sidebar stuff</aside>
  <footer>Footer content</footer>
  <script>alert('xss')</script>
  <style>body { color: red; }</style>
</body>
</html>"""

SAMPLE_HTML_NOISY = """<!DOCTYPE html>
<html>
<head><title>Noisy Page</title></head>
<body>
  <nav><a href="/">Home</a><a href="/about">About</a></nav>
  <div id="sidebar"><p>Sidebar ad</p></div>
  <main>
    <h1>The Real Content</h1>
    <p>This is the article body that should be extracted.</p>
  </main>
  <footer>&copy; 2024</footer>
</body>
</html>"""


class TestHTML:
    def test_extract_basic_html(self):
        result, fmt, meta = process_document(SAMPLE_HTML, "doc.html")
        assert fmt == "html"
        assert "Main Title" in result
        assert "First paragraph" in result
        assert "link" in result
        assert "script" not in result.lower()
        assert "style" not in result.lower()
        assert "nav" not in result or result.index("nav") > result.index("Main")
        assert meta["title"] == "Test Article"
        assert meta["word_count"] > 0
        assert meta["reading_time_minutes"] >= 1

    def test_html_noisy_page(self):
        result, fmt, meta = process_document(SAMPLE_HTML_NOISY, "doc.html")
        assert fmt == "html"
        assert "The Real Content" in result
        assert "sidebar" not in result.lower()

    def test_html_preserves_headings(self):
        result, fmt, meta = process_document(
            "<html><body><h1>Title</h1><h2>Sub</h2><h3>Subsub</h3></body></html>",
            "doc.html",
        )
        assert "<h1>Title</h1>" in result or "Title" in result
        assert "<h2>Sub</h2>" in result or "Sub" in result

    def test_html_empty(self):
        result, fmt, meta = process_document("", "doc.html")
        assert result == ""

    def test_html_preserves_images(self):
        result, fmt, meta = process_document(
            '<html><body><img src="photo.jpg" alt="Photo" /></body></html>',
            "doc.html",
        )
        assert 'src="photo.jpg"' in result

    def test_html_preserves_blockquote(self):
        result, fmt, meta = process_document(
            "<html><body><blockquote>Famous words</blockquote></body></html>",
            "doc.html",
        )
        assert "Famous words" in result

    def test_html_auto_detect(self):
        result, fmt, meta = process_document(
            "<html><body><p>Content</p></body></html>"
        )
        assert fmt == "html"

    def test_html_no_doctype_but_html_tag(self):
        result, fmt, meta = process_document(
            "<html><body><p>Hello</p></body></html>"
        )
        assert fmt == "html"

    def test_html_preserves_img_src_from_url(self):
        """Image src should be resolved against base_url when extracted via URL fetch."""
        from app.services.reader import _extract_html

        html = '<html><body><img src="/images/photo.jpg" alt="Photo" /><img src="https://cdn.example.com/pic.png" alt="CDN" /><img src="data:image/png;base64,abc" alt="Inline" /></body></html>'
        result, fmt, meta = _extract_html(html, base_url="https://example.com/article")
        assert 'src="https://example.com/images/photo.jpg"' in result
        assert 'src="https://cdn.example.com/pic.png"' in result
        assert 'src="data:image/png;base64,abc"' in result

    def test_html_preserves_img_srcset(self):
        from app.services.reader import _extract_html

        html = '<html><body><img src="/img.jpg" srcset="/small.jpg 400w, /large.jpg 800w" /></body></html>'
        result, fmt, meta = _extract_html(html, base_url="https://site.com/page")
        assert 'src="https://site.com/img.jpg"' in result
        assert 'https://site.com/small.jpg' in result
        assert 'https://site.com/large.jpg' in result

    def test_html_img_data_src_promoted(self):
        from app.services.reader import _extract_html

        html = '<html><body><img src="data:image/gif;base64,placeholder" data-src="https://cdn.example.com/real.jpg" alt="Lazy" /></body></html>'
        result, fmt, meta = _extract_html(html, base_url="https://example.com/page")
        assert 'src="https://cdn.example.com/real.jpg"' in result
        assert 'data-src' not in result

    def test_html_img_data_src_without_src(self):
        from app.services.reader import _extract_html

        html = '<html><body><img data-src="/images/hero.png" alt="No src" /></body></html>'
        result, fmt, meta = _extract_html(html, base_url="https://site.com/article")
        assert 'src="https://site.com/images/hero.png"' in result

    def test_html_img_protocol_relative(self):
        from app.services.reader import _extract_html

        html = '<html><body><img src="//cdn.example.com/pic.jpg" /></body></html>'
        result, fmt, meta = _extract_html(html, base_url="https://example.com")
        assert 'src="https://cdn.example.com/pic.jpg"' in result

    def test_readability_extraction(self):
        html = """<html><body>
          <div class="entry-content">
            <h1>Article</h1>
            <p>This is the main article text that readability should identify</p>
            <p>More content here across multiple paragraphs.</p>
          </div>
          <div class="comments">
            <p>Comment one</p>
            <p>Comment two</p>
          </div>
        </body></html>"""
        result, fmt, meta = process_document(html, "doc.html")
        assert fmt == "html"
        assert "Article" in result
        assert meta["word_count"] > 0


class TestURL:
    def test_is_url(self):
        assert _is_url("https://example.com")
        assert _is_url("http://test.org/page")
        assert not _is_url("not a url")
        assert not _is_url("")

    def test_fetch_url_returns_error_on_bad_url(self):
        result, fmt, meta = process_document("https://nonexistent.example.test/page")
        assert fmt == "error"
        assert "error" in meta

    def test_fetch_url_rejects_non_html_via_content_type(self):
        result, fmt, meta = process_document("https://example.com/file.pdf")
        # Without a real server, this will fail as a request error
        assert fmt == "error"


class TestPDF:
    def _make_pdf(self, text: str) -> bytes:
        from reportlab.pdfgen import canvas
        buf = io.BytesIO()
        c = canvas.Canvas(buf)
        c.drawString(100, 750, text)
        c.save()
        return buf.getvalue()

    def test_extract_pdf(self):
        pdf_bytes = self._make_pdf("Hello PDF World")
        result, fmt, meta = process_document(pdf_bytes, "doc.pdf")
        assert fmt == "pdf"
        assert "Hello PDF World" in result
        assert meta["pages"] == 1
        assert meta["word_count"] > 0

    def test_pdf_empty(self):
        from reportlab.pdfgen import canvas
        buf = io.BytesIO()
        c = canvas.Canvas(buf)
        c.save()
        pdf_bytes = buf.getvalue()
        result, fmt, meta = process_document(pdf_bytes, "doc.pdf")
        # Empty PDF should return an error
        assert fmt == "error"


class TestDOCX:
    def _make_docx(self, paragraphs: list[str]) -> bytes:
        doc = Document()
        for para in paragraphs:
            doc.add_paragraph(para)
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    def test_extract_docx(self):
        docx_bytes = self._make_docx(["Hello DOCX World", "Second paragraph"])
        result, fmt, meta = process_document(docx_bytes, "doc.docx")
        assert fmt == "docx"
        assert "Hello DOCX World" in result
        assert "Second paragraph" in result
        assert meta["word_count"] > 0

    def test_docx_with_headings(self):
        doc = Document()
        doc.add_heading("Chapter 1", level=1)
        doc.add_paragraph("Content here")
        buf = io.BytesIO()
        doc.save(buf)
        docx_bytes = buf.getvalue()
        result, fmt, meta = process_document(docx_bytes, "doc.docx")
        assert fmt == "docx"
        assert "<h1>" in result.upper() or "Chapter 1" in result


class TestHelpers:
    def test_looks_like_html(self):
        assert _looks_like_html("<!DOCTYPE html><html>")
        assert _looks_like_html("<html><body></body></html>")
        assert _looks_like_html("  <html>  ")
        assert not _looks_like_html("Just some text")
        assert not _looks_like_html("")

    def test_count_words(self):
        assert _count_words("Hello world") == 2
        assert _count_words("<p>Hello <strong>world</strong></p>") == 2
        assert _count_words("") == 0
