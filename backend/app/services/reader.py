import io
import re
from urllib.parse import urljoin

import httpx
from bs4 import BeautifulSoup, NavigableString, Tag  # type: ignore[attr-defined]
from pypdf import PdfReader
from readability import Document as ReadabilityDoc


def process_document(
    content: str | bytes,
    filename: str | None = None,
) -> tuple[str, str, dict[str, object]]:
    content_bytes: bytes
    if isinstance(content, str):
        content_bytes = content.encode("utf-8")
    else:
        content_bytes = content

    if _is_url(content_str(content_bytes)):
        return _fetch_url(content_str(content_bytes))

    ext = _get_ext(filename)

    if ext in ("pdf",):
        return _extract_pdf(content_bytes)

    if ext in ("docx",):
        return _extract_docx(content_bytes)

    text = content_str(content_bytes)

    if ext == "html" or _looks_like_html(text):
        return _extract_html(text)

    return _wrap_text(text, "text", {})


def content_str(content: bytes) -> str:
    try:
        return content.decode("utf-8")
    except UnicodeDecodeError:
        return content.decode("latin-1")


def _get_ext(filename: str | None) -> str:
    if not filename:
        return ""
    return filename.rsplit(".", 1)[-1].lower()


def _is_url(text: str) -> bool:
    return bool(re.match(r"^https?://", text.strip()))


def _looks_like_html(text: str) -> bool:
    head = text[:1000].strip()
    return bool(
        re.search(r"<!DOCTYPE html|<\s*html[\s>]", head, re.IGNORECASE)
    )


def _fetch_url(url: str) -> tuple[str, str, dict[str, object]]:
    import random

    user_agents = [
        "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36",
        "Mozilla/5.0 (Macintosh; Intel Mac OS X 14_2) AppleWebKit/605.1.15 (KHTML, like Gecko) Version/17.2 Safari/605.1.15",
        "Mozilla/5.0 (X11; Linux x86_64; rv:121.0) Gecko/20100101 Firefox/121.0",
    ]

    for attempt in range(2):
        try:
            with httpx.Client(
                timeout=httpx.Timeout(20.0, connect=10.0, read=15.0),
                follow_redirects=True,
            ) as client:
                resp = client.get(
                    url,
                    headers={
                        "User-Agent": random.choice(user_agents),
                        "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
                        "Accept-Language": "en-US,en;q=0.5",
                    },
                )
                resp.raise_for_status()
                content_type = resp.headers.get("content-type", "").lower()
                if "text/html" not in content_type and "application/xhtml" not in content_type:
                    if "application/pdf" in content_type:
                        return _extract_pdf(resp.content)
                    return _wrap_text(
                        resp.text[:500],
                        "unsupported",
                        {"error": f"Unsupported content type: {content_type}", "url": url},
                    )
                html = resp.text
                final_url = str(resp.url)

                result, fmt, meta = _extract_html(html, base_url=final_url)

                if not result or len(result) < 50:
                    result, fmt, meta = _extract_html_fallback(html, base_url=final_url)

                meta["url"] = final_url
                meta["status_code"] = resp.status_code
                if attempt > 0:
                    meta["retry"] = True
                return result, fmt, meta
        except httpx.TimeoutException:
            if attempt == 0:
                continue
            return _wrap_text("", "error", {"error": "Request timed out after 20 seconds", "url": url})
        except httpx.HTTPStatusError as e:
            status = e.response.status_code
            if status in (429, 503, 502) and attempt == 0:
                continue
            return _wrap_text("", "error", {"error": f"HTTP {status}", "url": url})
        except httpx.RequestError as e:
            if attempt == 0:
                continue
            return _wrap_text("", "error", {"error": f"Request failed: {e!s}", "url": url})

    return _wrap_text("", "error", {"error": "Failed after retries", "url": url})


def _extract_html_fallback(
    html: str,
    base_url: str | None = None,
) -> tuple[str, str, dict[str, object]]:
    soup = BeautifulSoup(html, "lxml")
    for tag in soup.find_all(["script", "style", "nav", "footer", "header", "aside", "noscript", "iframe", "form", "svg"]):
        tag.decompose()

    body = soup.find("body") or soup

    for tag in body.find_all(True):
        if isinstance(tag, Tag):
            tag.attrs = {}

    allowed = {"h1", "h2", "h3", "h4", "h5", "h6", "p", "li", "pre", "blockquote", "div", "article", "section", "ul", "ol", "br", "hr", "table", "tr", "td", "th"}
    for tag in body.find_all(True):
        if isinstance(tag, Tag) and tag.name not in allowed:
            tag.unwrap()

    html_out = ""
    for elem in body.children:
        if isinstance(elem, Tag):
            text = elem.get_text(strip=True)
            if text and len(text) > 20:
                html_out += f"<p>{_escape_html(text)}</p>\n"

    word_count = _count_words(html_out)
    reading_time = max(1, round(word_count / 200))
    meta: dict[str, object] = {
        "word_count": word_count,
        "reading_time_minutes": reading_time,
        "source_type": "html_fallback",
    }
    return html_out, "html", meta


def _extract_html(
    html: str,
    base_url: str | None = None,
) -> tuple[str, str, dict[str, object]]:
    if not html.strip():
        return "", "html", {"word_count": 0, "reading_time_minutes": 0, "source_type": "html"}

    doc = ReadabilityDoc(html)
    title = doc.title() or ""
    summary_html = doc.summary()

    soup = BeautifulSoup(summary_html, "lxml")

    _clean_soup(soup)

    if base_url:
        _resolve_images(soup, base_url)

    body = soup.find("body")
    root = body if body else soup

    result = _soup_to_html(root)
    word_count = _count_words(result)
    reading_time = max(1, round(word_count / 200))

    meta: dict[str, object] = {
        "title": title,
        "word_count": word_count,
        "reading_time_minutes": reading_time,
        "source_type": "html",
    }
    if title:
        meta["title"] = title

    return result, "html", meta


def _clean_soup(soup: BeautifulSoup) -> None:
    for tag in soup.find_all(
        ["script", "style", "nav", "footer", "header", "aside", "noscript", "iframe", "form"]
    ):
        tag.decompose()

    for tag in soup.find_all(True):
        if isinstance(tag, Tag):
            allowed = {"a", "img", "picture", "source", "figure", "figcaption", "abbr", "acronym"}
            attrs_to_keep = {"href", "src", "srcset", "data-src", "media", "type", "alt", "title"}
            if tag.name in allowed:
                tag.attrs = {k: v for k, v in tag.attrs.items() if k in attrs_to_keep}
            else:
                tag.attrs = {}

    for tag in soup.find_all(["b", "strong"]):
        tag.name = "strong"
    for tag in soup.find_all(["i", "em"]):
        tag.name = "em"

    allowed_names = {
        "a", "abbr", "acronym", "b", "big", "blockquote", "br", "caption",
        "center", "cite", "code", "col", "colgroup", "dd", "del", "dfn",
        "dir", "div", "dl", "dt", "em", "figure", "font", "h1", "h2",
        "h3", "h4", "h5", "h6", "hr", "i", "img", "ins", "kbd", "label",
        "legend", "li", "menu", "ol", "p", "pre", "q", "s", "samp",
        "small", "span", "strike", "strong", "sub", "sup", "table",
        "tbody", "td", "tfoot", "th", "thead", "tr", "tt", "u", "ul",
        "var",
    }
    for elem in soup.find_all(True):
        if isinstance(elem, Tag) and elem.name not in allowed_names:
            elem.unwrap()

    for tag in soup.find_all(True):
        if isinstance(tag, Tag):
            content_text = tag.get_text(strip=True)
            if not content_text and tag.name not in ("br", "hr", "img") and not tag.find("img"):
                tag.decompose()

    for a in soup.find_all("a"):
        if isinstance(a, Tag):
            children = list(a.children)
            if len(children) == 1 and isinstance(children[0], Tag) and children[0].name == "img":
                a.unwrap()

    for tag in soup.find_all(["blockquote", "pre", "code"]):
        if isinstance(tag, Tag):
            inner = tag.get_text(strip=True)
            if not inner:
                tag.decompose()
                continue
            tag.clear()
            tag.append(inner)


def _resolve_images(soup: BeautifulSoup, base_url: str) -> None:
    for img in soup.find_all("img"):
        src_raw = img.get("src")
        data_src_raw = img.get("data-src")

        src = str(src_raw) if src_raw else ""
        data_src = str(data_src_raw) if data_src_raw else ""

        if data_src and (not src or src.startswith("data:")):
            img["src"] = data_src
            if "data-src" in img.attrs:
                del img["data-src"]
            src = data_src

        if src and not src.startswith(("http://", "https://", "data:", "//")):
            img["src"] = urljoin(base_url, src)
        elif src and src.startswith("//"):
            img["src"] = "https:" + src

        _resolve_attr_srcset(img, base_url, "srcset")

    for source in soup.find_all("source"):
        _resolve_attr_srcset(source, base_url, "srcset")
        src_raw = source.get("src")
        src_val = str(src_raw) if src_raw else ""
        if src_val and not src_val.startswith(("http://", "https://", "data:", "//")):
            source["src"] = urljoin(base_url, src_val)
        elif src_val and src_val.startswith("//"):
            source["src"] = "https:" + src_val


def _resolve_attr_srcset(tag: Tag, base_url: str, attr: str) -> None:
    raw = tag.get(attr)
    if not raw:
        return
    val = str(raw)
    resolved: list[str] = []
    for entry in val.split(","):
        entry = entry.strip()
        if not entry:
            continue
        parts = entry.split()
        url_part = parts[0]
        if url_part.startswith("//"):
            parts[0] = "https:" + url_part
        elif not url_part.startswith(("http://", "https://", "data:")):
            parts[0] = urljoin(base_url, url_part)
        resolved.append(" ".join(parts))
    tag[attr] = ", ".join(resolved)


def _soup_to_html(root: Tag) -> str:
    result = ""
    for child in root.children:
        if isinstance(child, NavigableString):
            text = str(child).strip()
            if text:
                result += f"<p>{_escape_html(text)}</p>\n"
        elif isinstance(child, Tag):
            result += str(child) + "\n"
    return result.strip()


def _escape_html(text: str) -> str:
    return (
        text.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
    )


def _extract_pdf(content: bytes) -> tuple[str, str, dict[str, object]]:
    try:
        reader = PdfReader(io.BytesIO(content))
    except Exception:
        return _wrap_text(
            "", "error",
            {"error": "Could not read PDF (password protected or corrupted)"},
        )

    num_pages = len(reader.pages)
    if num_pages == 0:
        return _wrap_text("", "error", {"error": "PDF has no pages"})

    paragraphs: list[str] = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            for para in text.split("\n"):
                para = para.strip()
                if para:
                    paragraphs.append(para)

    if not paragraphs:
        return _wrap_text(
            "",
            "pdf",
            {
                "error": "PDF appears to be a scanned document (no extractable text)",
                "pages": num_pages,
                "source_type": "pdf",
            },
        )

    full_text = "\n\n".join(paragraphs)
    word_count = _count_words(full_text)
    reading_time = max(1, round(word_count / 200))

    html = "<p>" + "</p>\n<p>".join(_escape_html(p) for p in paragraphs) + "</p>"

    word_count = _count_words(html)
    reading_time = max(1, round(word_count / 200))

    meta: dict[str, object] = {
        "word_count": word_count,
        "reading_time_minutes": reading_time,
        "pages": num_pages,
        "source_type": "pdf",
    }

    return html, "pdf", meta


def _extract_docx(content: bytes) -> tuple[str, str, dict[str, object]]:
    try:
        from docx import Document
    except ImportError:
        return _wrap_text("", "error", {"error": "DOCX support not available"})

    try:
        doc = Document(io.BytesIO(content))
    except Exception:
        return _wrap_text("", "error", {"error": "Could not read DOCX file"})

    html_parts: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style = para.style.name.lower().replace(" ", "") if para.style else ""
        if style.startswith("heading1") or style.startswith("heading 1"):
            html_parts.append(f"<h1>{_escape_html(text)}</h1>")
        elif style.startswith("heading2") or style.startswith("heading 2"):
            html_parts.append(f"<h2>{_escape_html(text)}</h2>")
        elif style.startswith("heading3") or style.startswith("heading 3"):
            html_parts.append(f"<h3>{_escape_html(text)}</h3>")
        else:
            html_parts.append(f"<p>{_escape_html(text)}</p>")

    for table in doc.tables:
        html_parts.append("<table>")
        for row in table.rows:
            html_parts.append("<tr>")
            for cell in row.cells:
                html_parts.append(f"<td>{_escape_html(cell.text.strip())}</td>")
            html_parts.append("</tr>")
        html_parts.append("</table>")

    result = "\n".join(html_parts)

    if not html_parts:
        err_meta: dict[str, object] = {
            "error": "No extractable content in DOCX", "source_type": "docx",
        }
        return _wrap_text("", "docx", err_meta)

    word_count = _count_words(result)
    reading_time = max(1, round(word_count / 200))

    meta: dict[str, object] = {
        "word_count": word_count,
        "reading_time_minutes": reading_time,
        "source_type": "docx",
    }

    return result, "docx", meta


def _wrap_text(text: str, fmt: str, extra: dict[str, object]) -> tuple[str, str, dict[str, object]]:
    return text, fmt, extra


def _count_words(text: str) -> int:
    text = re.sub(r"<[^>]+>", " ", text)
    return len(re.findall(r"\b\w+\b", text))
